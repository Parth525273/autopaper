"""
DOCX Generation Engine
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from app.core.templates import get_template


def strip_html(html: str) -> str:
    html = html.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"')
    clean = re.sub(r"<[^>]+>", " ", html)
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def parse_html_to_blocks(html: str) -> list:
    if not html:
        return []

    html = html.replace("&nbsp;", " ").replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    blocks = []
    list_blocks = {}
    counter = [0]
    li_pattern = re.compile(r'<li[^>]*>(.*?)</li>', re.DOTALL)

    def replace_ul(m):
        key = f"__LIST_{counter[0]}__"
        items = li_pattern.findall(m.group(1))
        list_blocks[key] = [{"type": "bullet", "text": re.sub(r"<[^>]+>", "", i).strip()} for i in items if re.sub(r"<[^>]+>", "", i).strip()]
        counter[0] += 1
        return key

    def replace_ol(m):
        key = f"__LIST_{counter[0]}__"
        items = li_pattern.findall(m.group(1))
        list_blocks[key] = [{"type": "ordered", "text": re.sub(r"<[^>]+>", "", i).strip(), "index": idx + 1} for idx, i in enumerate(items) if re.sub(r"<[^>]+>", "", i).strip()]
        counter[0] += 1
        return key

    processed = re.sub(r'<ul>(.*?)</ul>', replace_ul, html, flags=re.DOTALL)
    processed = re.sub(r'<ol>(.*?)</ol>', replace_ol, processed, flags=re.DOTALL)
    parts = re.split(r'(__LIST_\d+__)', processed)

    para_pattern = re.compile(r'<p[^>]*>(.*?)</p>', re.DOTALL)
    h_pattern = re.compile(r'<h([1-4])[^>]*>(.*?)</h[1-4]>', re.DOTALL)

    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part in list_blocks:
            blocks.extend(list_blocks[part])
        else:
            all_elements = []
            for m in para_pattern.finditer(part):
                all_elements.append(("p", m.start(), m.group(1)))
            for m in h_pattern.finditer(part):
                all_elements.append((f"h{m.group(1)}", m.start(), m.group(2)))
            all_elements.sort(key=lambda x: x[1])

            for el_type, _, inner in all_elements:
                text = re.sub(r"<[^>]+>", "", inner).strip()
                if not text:
                    continue
                if el_type == "p":
                    blocks.append({"type": "paragraph", "text": text, "bold": "<strong>" in inner, "italic": "<em>" in inner})
                elif el_type.startswith("h"):
                    blocks.append({"type": "heading", "text": text, "level": int(el_type[1])})

    return blocks


def parse_references(html: str) -> list:
    """Parse references section - each line/paragraph is one reference."""
    if not html:
        return []
    html = html.replace("&nbsp;", " ").replace("&amp;", "&")
    para_pattern = re.compile(r'<p[^>]*>(.*?)</p>', re.DOTALL)
    refs = []
    for m in para_pattern.finditer(html):
        text = re.sub(r"<[^>]+>", "", m.group(1)).strip()
        if text:
            refs.append(text)
    return refs


def add_run_with_font(paragraph, text: str, font_name: str, font_size: int,
                      bold=False, italic=False, all_caps=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.all_caps = all_caps
    return run


def set_two_columns(section):
    sectPr = section._sectPr
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")
    sectPr.append(cols)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_ieee_authors_table(doc, authors: list, font_name: str):
    if not authors:
        return
    raw_rows = [authors[i:i+3] for i in range(0, len(authors), 3)]
    rows = [row + [None] * (3 - len(row)) for row in raw_rows]

    for row_authors in rows:
        table = doc.add_table(rows=1, cols=3)
        remove_table_borders(table)
        for i, author in enumerate(row_authors):
            cell = table.rows[0].cells[i]
            cell.width = Inches(2.4)
            cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
            if author is None:
                cell.add_paragraph()
                continue

            def add_cell_para(text, bold=False, italic=False, size=10, _cell=cell):
                p = _cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                if text:
                    add_run_with_font(p, text, font_name, size, bold=bold, italic=italic)

            add_cell_para(author.get("name", ""), size=11)
            if author.get("department"):
                add_cell_para(author["department"], italic=True, size=10)
            if author.get("affiliation"):
                add_cell_para(author["affiliation"], italic=True, size=10)
            if author.get("email"):
                add_cell_para(author["email"], size=10)
        doc.add_paragraph()


def add_content_blocks(doc, blocks: list, body_font: dict, spacing: dict):
    for block in blocks:
        if block["type"] == "paragraph":
            text = block["text"]
            if not text.strip():
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(spacing.get("paragraph_space_after_pt", 6))
            add_run_with_font(p, text, body_font["name"], body_font["size_pt"],
                              bold=block.get("bold", False), italic=block.get("italic", False))

        elif block["type"] == "heading":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            add_run_with_font(p, block["text"], body_font["name"], body_font["size_pt"],
                              bold=True, italic=block["level"] >= 3)

        elif block["type"] == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_run_with_font(p, block["text"], body_font["name"], body_font["size_pt"])

        elif block["type"] == "ordered":
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            add_run_with_font(p, block["text"], body_font["name"], body_font["size_pt"])


def add_ieee_references(doc, refs_html: str, body_font: dict, ref_font: dict):
    """Add references in IEEE format: [1] Author, Title, Journal, year."""
    refs = parse_references(refs_html)
    for idx, ref_text in enumerate(refs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(4)

        # Add [n] number in normal weight
        num_run = p.add_run(f"[{idx + 1}]  ")
        num_run.font.name = ref_font["name"]
        num_run.font.size = Pt(ref_font["size_pt"])
        num_run.font.bold = False

        # Add reference text
        text_run = p.add_run(ref_text)
        text_run.font.name = ref_font["name"]
        text_run.font.size = Pt(ref_font["size_pt"])
        text_run.font.bold = False


def generate_docx(paper_data: dict, output_path: str) -> str:
    template_id = paper_data.get("templateId", "ieee-conference")
    template = get_template(template_id)

    doc = Document()
    page_cfg = template["page"]
    fonts = template["fonts"]
    spacing = template["spacing"]
    numbering = template["numbering"]
    is_ieee = template_id in ["ieee-conference", "ieee-journal"]

    section = doc.sections[0]
    section.page_width = Inches(page_cfg["width_inches"])
    section.page_height = Inches(page_cfg["height_inches"])
    section.top_margin = Inches(page_cfg["margin_top_inches"])
    section.bottom_margin = Inches(page_cfg["margin_bottom_inches"])
    section.left_margin = Inches(page_cfg["margin_left_inches"])
    section.right_margin = Inches(page_cfg["margin_right_inches"])

    title_font = fonts["title"]
    body_font = fonts["body"]
    heading_font = fonts["section_heading"]
    abstract_heading_font = fonts["abstract_heading"]
    abstract_body_font = fonts["abstract_body"]
    keyword_font = fonts["keywords"]
    ref_font = fonts["references"]

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    add_run_with_font(title_para, paper_data.get("title", "Untitled Paper"),
                      title_font["name"], title_font["size_pt"], bold=True)

    # Authors
    authors = paper_data.get("authors", [])
    if is_ieee and authors:
        add_ieee_authors_table(doc, authors, fonts["authors"]["name"])
    else:
        author_names = ", ".join([a.get("name", "") for a in authors if a.get("name")])
        if author_names:
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_after = Pt(4)
            add_run_with_font(ap, author_names, fonts["authors"]["name"], fonts["authors"]["size_pt"])
        for author in authors:
            parts = []
            if author.get("department"): parts.append(author["department"])
            if author.get("affiliation"): parts.append(author["affiliation"])
            if parts:
                ap2 = doc.add_paragraph()
                ap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ap2.paragraph_format.space_after = Pt(2)
                add_run_with_font(ap2, ", ".join(parts), fonts["affiliation"]["name"],
                                  fonts["affiliation"].get("size_pt", 10),
                                  italic=fonts["affiliation"].get("italic", False))
            if author.get("email"):
                ep = doc.add_paragraph()
                ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ep.paragraph_format.space_after = Pt(2)
                add_run_with_font(ep, author["email"], fonts["authors"]["name"], 10)

    # Two-column break
    if template["columns"] == 2:
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_two_columns(new_section)
        new_section.top_margin = Inches(page_cfg["margin_top_inches"])
        new_section.bottom_margin = Inches(page_cfg["margin_bottom_inches"])
        new_section.left_margin = Inches(page_cfg["margin_left_inches"])
        new_section.right_margin = Inches(page_cfg["margin_right_inches"])

    # Abstract
    abstract_text = strip_html(paper_data.get("sections", {}).get("abstract", ""))
    if abstract_text:
        abs_para = doc.add_paragraph()
        abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        abs_para.paragraph_format.space_after = Pt(4)
        add_run_with_font(abs_para, "Abstract\u2014", abstract_heading_font["name"],
                          abstract_heading_font.get("size_pt", 9), bold=True, italic=True)
        add_run_with_font(abs_para, " " + abstract_text, abstract_body_font["name"],
                          abstract_body_font.get("size_pt", 9), bold=True)

    # Keywords
    keywords = paper_data.get("keywords", [])
    if keywords:
        kw_para = doc.add_paragraph()
        kw_para.paragraph_format.space_after = Pt(8)
        add_run_with_font(kw_para, "Keywords\u2014 " + ", ".join(keywords),
                          keyword_font["name"], keyword_font.get("size_pt", 9), bold=True, italic=True)

    # ── FIXED section order: Future Scope before Conclusion ──
    section_order = [
        ("introduction",   "Introduction"),
        ("literatureReview", "Literature Review"),
        ("methodology",    "Methodology"),
        ("systemDesign",   "System Design"),
        ("implementation", "Implementation"),
        ("results",        "Results and Discussion"),
        ("futureScope",    "Future Scope"),
        ("conclusion",     "Conclusion"),
        ("acknowledgement","Acknowledgement"),
        ("references",     "References"),
    ]

    roman_numerals = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI"]
    section_count = 0

    for key, default_label in section_order:
        raw_html = paper_data.get("sections", {}).get(key, "")
        if not raw_html or not strip_html(raw_html):
            continue

        if numbering["sections"]:
            label = f"{roman_numerals[section_count]}. {default_label.upper()}" if is_ieee else f"{section_count + 1}. {default_label}"
        else:
            label = default_label.upper()

        heading_para = doc.add_paragraph()
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_para.paragraph_format.space_before = Pt(spacing["section_space_before_pt"])
        heading_para.paragraph_format.space_after = Pt(4)
        add_run_with_font(heading_para, label, heading_font["name"], heading_font["size_pt"],
                          bold=heading_font.get("bold", True), all_caps=heading_font.get("all_caps", False))

        # References get special IEEE formatting
        if key == "references" and is_ieee:
            add_ieee_references(doc, raw_html, body_font, ref_font)
        else:
            blocks = parse_html_to_blocks(raw_html)
            add_content_blocks(doc, blocks, body_font, spacing)

        section_count += 1

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path
